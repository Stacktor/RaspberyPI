from docx import Document


# Neues Dokument erstellen mit Deckblatt und Inhaltsverzeichnis
doc = Document()

# Deckblatt
doc.add_heading('Wasserstandsüberwachung mit Raspberry Pi 4 und HC-SR04 Ultraschallsensor', 0)
doc.add_paragraph("Projektbericht")
doc.add_paragraph("")
doc.add_paragraph("Autor: [Dein Name]")
doc.add_paragraph("Datum: [Datum]")
doc.add_page_break()

# Inhaltsverzeichnis
doc.add_heading('Inhaltsverzeichnis', level=1)
doc.add_paragraph('1. Einleitung', style='List Number')
doc.add_paragraph('2. Hardware-Komponenten', style='List Number')
doc.add_paragraph('3. Software-Code', style='List Number')
doc.add_paragraph('   3.1. Grundkonfiguration', style='List Number 2')
doc.add_paragraph('   3.2. Modulkonfiguration', style='List Number 2')
doc.add_paragraph('   3.3. Code zur Messung des Wasserstands', style='List Number 2')
doc.add_paragraph('   3.4. Code zur Steuerung der Philips Hue Steckdose', style='List Number 2')
doc.add_paragraph('   3.5. Integrierter Code zur Überwachung und Steuerung', style='List Number 2')
doc.add_paragraph('4. Hardware-Liste', style='List Number')
doc.add_paragraph('5. Blockschaltbild', style='List Number')
doc.add_paragraph('6. Projekt-Ablauf', style='List Number')
doc.add_paragraph('7. Fazit', style='List Number')
doc.add_page_break()

# Einleitung
doc.add_heading('1. Einleitung', level=1)
doc.add_paragraph(
    "Dieses Projekt beschreibt die Realisierung einer Wasserstandsüberwachung für einen Eimer mittels eines Raspberry Pi 4 und eines HC-SR04 Ultraschallsensors. "
    "Sobald der Wasserstand 95% erreicht, wird eine Philips Hue Steckdose angesteuert, die eine Pumpe aktiviert, um den Eimer zu entleeren."
)

# Hardware-Komponenten
doc.add_heading('2. Hardware-Komponenten', level=1)
doc.add_paragraph(
    "- Raspberry Pi 4\n"
    "- HC-SR04 Ultraschallsensor\n"
    "- Philips Hue Steckdose\n"
    "- Wasserpumpe\n"
    "- Eimer\n"
    "- Diverse Kabel\n"
    "- Steckbrett (Breadboard)"
)

# Beispielbild hinzufügen (Platzhalter für Hardware-Komponenten)
doc.add_paragraph("![Hardware-Komponenten](./bilder/hardware.jpg)")

# Software-Code
doc.add_heading('3. Software-Code', level=1)

doc.add_heading('Grundkonfiguration', level=2)
doc.add_paragraph(
    "Der Raspberry Pi wird mit dem Betriebssystem Raspbian ausgestattet und alle notwendigen Pakete werden installiert.\n"
    "```bash\n"
    "sudo apt update\n"
    "sudo apt install python3 python3-pip\n"
    "pip3 install phue\n"
    "```"
)

doc.add_heading('Modulkonfiguration', level=2)
doc.add_paragraph(
    "Der HC-SR04 Sensor wird über die GPIO Pins des Raspberry Pi ausgelesen. Hierzu wird eine Python-Bibliothek verwendet, um die Entfernungsdaten zu messen und den Wasserstand zu berechnen."
)

doc.add_heading('Code zur Messung des Wasserstands', level=3)
doc.add_paragraph(
    "```python\n"
    "import RPi.GPIO as GPIO\n"
    "import time\n\n"
    "# Pin-Definitionen\n"
    "TRIG = 23\n"
    "ECHO = 24\n\n"
    "GPIO.setmode(GPIO.BCM)\n"
    "GPIO.setup(TRIG, GPIO.OUT)\n"
    "GPIO.setup(ECHO, GPIO.IN)\n\n"
    "def get_distance():\n"
    "    GPIO.output(TRIG, True)\n"
    "    time.sleep(0.00001)\n"
    "    GPIO.output(TRIG, False)\n\n"
    "    while GPIO.input(ECHO) == 0:\n"
    "        start_time = time.time()\n\n"
    "    while GPIO.input(ECHO) == 1:\n"
    "        end_time = time.time()\n\n"
    "    duration = end_time - start_time\n"
    "    distance = (duration * 34300) / 2\n\n"
    "    return distance\n\n"
    "try:\n"
    "    while True:\n"
    "        distance = get_distance()\n"
    "        print(f\"Distance: {distance:.2f} cm\")\n"
    "        time.sleep(1)\n"
    "except KeyboardInterrupt:\n"
    "    GPIO.cleanup()\n"
    "```"
)

doc.add_heading('Code zur Steuerung der Philips Hue Steckdose', level=3)
doc.add_paragraph(
    "```python\n"
    "from phue import Bridge\n\n"
    "bridge_ip = 'IP_ADDRESS_OF_BRIDGE'\n"
    "b = Bridge(bridge_ip)\n"
    "b.connect()\n\n"
    "def control_pump(turn_on):\n"
    "    lights = b.get_light_objects('name')\n"
    "    pump_light = lights['PumpSocket']\n\n"
    "    if turn_on:\n"
    "        pump_light.on = True\n"
    "    else:\n"
    "        pump_light.on = False\n\n"
    "# Beispiel zur Steuerung der Pumpe\n"
    "control_pump(True)  # Pumpe einschalten\n"
    "control_pump(False)  # Pumpe ausschalten\n"
    "```"
)

doc.add_heading('Integrierter Code zur Überwachung und Steuerung', level=3)
doc.add_paragraph(
    "```python\n"
    "import RPi.GPIO as GPIO\n"
    "import time\n"
    "from phue import Bridge\n\n"
    "# Pin-Definitionen und Philips Hue Bridge\n"
    "TRIG = 23\n"
    "ECHO = 24\n"
    "bridge_ip = 'IP_ADDRESS_OF_BRIDGE'\n"
    "b = Bridge(bridge_ip)\n"
    "b.connect()\n\n"
    "GPIO.setmode(GPIO.BCM)\n"
    "GPIO.setup(TRIG, GPIO.OUT)\n"
    "GPIO.setup(ECHO, GPIO.IN)\n\n"
    "def get_distance():\n"
    "    GPIO.output(TRIG, True)\n"
    "    time.sleep(0.00001)\n"
    "    GPIO.output(TRIG, False)\n\n"
    "    while GPIO.input(ECHO) == 0:\n"
    "        start_time = time.time()\n\n"
    "    while GPIO.input(ECHO) == 1:\n"
    "        end_time = time.time()\n\n"
    "    duration = end_time - start_time\n"
    "    distance = (duration * 34300) / 2\n"
    "    return distance\n\n"
    "def control_pump(turn_on):\n"
    "    lights = b.get_light_objects('name')\n"
    "    pump_light = lights['PumpSocket']\n\n"
    "    if turn_on:\n"
    "        pump_light.on = True\n"
    "    else:\n"
    "        pump_light.on = False\n\n"
    "try:\n"
    "    while True:\n"
    "        distance = get_distance()\n"
    "        water_level_percentage = (MAX_WATER_LEVEL - distance) / MAX_WATER_LEVEL * 100\n\n"
    "        if water_level_percentage >= 95:\n"
    "            control_pump(True)\n"
    "        else:\n"
    "            control_pump(False)\n\n"
    "        print(f\"Water Level: {water_level_percentage:.2f}%\")\n"
    "        time.sleep(1)\n"
    "except KeyboardInterrupt:\n"
    "    GPIO.cleanup()\n"
    "```"
)

# Hardware-Liste
doc.add_heading('4. Hardware-Liste', level=1)
doc.add_paragraph(
    "- 1x Raspberry Pi 4\n"
    "- 1x HC-SR04 Ultraschallsensor\n"
    "- 1x Philips Hue Steckdose\n"
    "- 1x Wasserpumpe\n"
    "- 1x Eimer\n"
    "- 1x Steckbrett (Breadboard)\n"
    "- Verbindungskabel"
)

# Blockschaltbild
doc.add_heading('5. Blockschaltbild', level=1)
doc.add_paragraph("![Blockschaltbild](./bilder/blockschaltbild.jpg)")

# Projekt-Ablauf
doc.add_heading('6. Projekt-Ablauf', level=1)
doc.add_paragraph(
    "1. Aufbau der Hardware:\n"
    "   - Verkabelung des HC-SR04 Sensors mit dem Raspberry Pi.\n"
    "   - Verbindung der Wasserpumpe mit der Philips Hue Steckdose.\n\n"
    "2. Software-Installation:\n"
    "   - Installation des Betriebssystems auf dem Raspberry Pi.\n"
    "   - Installation der notwendigen Python-Pakete.\n\n"
    "3. Programmierung:\n"
    "   - Schreiben und Testen des Codes zur Messung des Wasserstands.\n"
    "   - Implementierung der Steuerung der Philips Hue Steckdose.\n\n"
    "4. Integration und Test:\n"
    "   - Zusammenführen des Codes.\n"
    "   - Testen des gesamten Systems im realen Szenario."
)

# Fazit
doc.add_heading('7. Fazit', level=1)
doc.add_paragraph(
    "Das Projekt zeigt eine einfache Methode zur Überwachung des Wasserstands und automatisierten Steuerung einer Pumpe mittels eines Raspberry Pi und eines HC-SR04 Ultraschallsensors. "
    "Durch die Nutzung der Philips Hue Steckdose kann die Steuerung bequem über das bestehende Smart-Home-System erfolgen. Das System ist zuverlässig und kann bei Bedarf erweitert werden."
)

# Dokument speichern
file_path = "Wasserstandsüberwachung_Projektbericht.docx"
doc.save(file_path)

file_path
